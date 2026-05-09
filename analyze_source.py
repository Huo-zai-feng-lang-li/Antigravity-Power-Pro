import docx
import json

def extract_functional_points(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # Analyze tables if present
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content.append({"type": "table", "data": table_data})
        
    # Analyze paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content.append({"type": "paragraphs", "data": paragraphs})
    
    return content

source_path = r"c:\Users\Administrator\Desktop\软件系统功能检查表.docx"
extracted_info = extract_functional_points(source_path)

with open("source_analysis.json", "w", encoding="utf-8") as f:
    json.dump(extracted_info, f, ensure_ascii=False, indent=2)

print("Source analysis complete.")
